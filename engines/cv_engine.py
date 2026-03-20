import os 
import json 
import shutil 
from docx import Document 
from docx.shared import Pt, RGBColor 
from datetime import date 
import sys 
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))) 
 
def load_master_cv(cv_path="master_cv.docx"): 
    doc = Document(cv_path) 
    full_text = [] 
    for para in doc.paragraphs: 
        if para.text.strip(): 
            full_text.append(para.text.strip()) 
    return "\n".join(full_text), doc 
 
def extract_cv_summary(cv_text): 
    lines = cv_text.split("\n") 
    summary_lines = [] 
    capture = False 
    for line in lines: 
        if "EXECUTIVE PROFILE" in line.upper(): 
            capture = True 
            continue 
        if capture and line.strip(): 
            if any(section in line.upper() for section in ["CORE LEADERSHIP", "PROFESSIONAL EXPERIENCE", "COMPETENCIES", "EDUCATION"]): 
                break 
            summary_lines.append(line) 
    return " ".join(summary_lines) 
 
def create_tailored_cv(job_id, company, role, tailored_summary, cover_letter, master_cv_path="master_cv.docx"): 
    from docx.shared import Pt 
    from copy import deepcopy 
    import re 
 
    os.makedirs("applications", exist_ok=True) 
    safe_company = re.sub(r'[^\w\s-]', '', company).strip().replace(' ', '_')[:30] 
    safe_role = re.sub(r'[^\w\s-]', '', role).strip().replace(' ', '_')[:25] 
    folder_name = f"applications/{safe_company}_{safe_role}_{date.today().strftime('%d%b%Y')}" 
    os.makedirs(folder_name, exist_ok=True) 
 
    # Open fresh copy of master CV 
    doc = Document(master_cv_path) 
 
    # Find and replace ONLY the summary paragraphs after EXECUTIVE PROFILE 
    profile_found = False 
    summary_replaced = False 
    paragraphs_to_clear = [] 
 
    for i, para in enumerate(doc.paragraphs): 
        if "EXECUTIVE PROFILE" in para.text.upper(): 
            profile_found = True 
            continue 
        if profile_found and not summary_replaced: 
            if para.text.strip(): 
                if any(section in para.text.upper() for section in ["CORE LEADERSHIP", "PROFESSIONAL", "COMPETENCIES", "EDUCATION"]): 
                    break 
                paragraphs_to_clear.append(i) 
 
    # Replace first summary paragraph with tailored version, clear the rest 
    for idx, para_idx in enumerate(paragraphs_to_clear): 
        para = doc.paragraphs[para_idx] 
        if idx == 0: 
            # Keep formatting, just replace text 
            for run in para.runs: 
                run.text = "" 
            if para.runs: 
                para.runs[0].text = tailored_summary 
            else: 
                run = para.add_run(tailored_summary) 
                run.font.size = Pt(10) 
            summary_replaced = True 
        else: 
            # Clear additional summary paragraphs 
            for run in para.runs: 
                run.text = "" 
            para.clear() 
 
    cv_filename = f"{folder_name}/CV_{safe_company}_{safe_role}.docx" 
    doc.save(cv_filename) 
 
    # Create cover letter as clean document 
    cl_doc = Document() 
    cl_doc.add_paragraph(f"{date.today().strftime('%d %B %Y')}") 
    cl_doc.add_paragraph("") 
    cl_doc.add_paragraph(f"Re: {role} — {company}") 
    cl_doc.add_paragraph("") 
    for para_text in cover_letter.split("\n\n"): 
        if para_text.strip(): 
            p = cl_doc.add_paragraph(para_text.strip()) 
            p.paragraph_format.space_after = Pt(12) 
    cl_doc.add_paragraph("") 
    cl_doc.add_paragraph("Yours sincerely,") 
    cl_doc.add_paragraph(f"Syed Danish Hasan") 
    cl_doc.add_paragraph("+91 998 107 3077 | dhasan111@gmail.com") 
 
    cl_filename = f"{folder_name}/CoverLetter_{safe_company}_{safe_role}.docx" 
    cl_doc.save(cl_filename) 
 
    jd_filename = f"{folder_name}/JobDetails.txt" 
    with open(jd_filename, 'w') as f: 
        f.write(f"Company: {company}\nRole: {role}\nDate: {date.today().isoformat()}\n") 
 
    print(f"✓ CV saved: {cv_filename}") 
    print(f"✓ Cover letter saved: {cl_filename}") 
    return cv_filename, cl_filename, folder_name 
 
def convert_docx_to_pdf(docx_path): 
    import subprocess 
    output_dir = os.path.dirname(docx_path) 
    try: 
        result = subprocess.run([ 
            '/Applications/LibreOffice.app/Contents/MacOS/soffice', 
            '--headless', '--convert-to', 'pdf', 
            '--outdir', output_dir, docx_path 
        ], capture_output=True, text=True, timeout=30) 
        pdf_path = docx_path.replace('.docx', '.pdf') 
        if os.path.exists(pdf_path): 
            print(f"✓ PDF created: {pdf_path}") 
            return pdf_path 
        else: 
            print(f"PDF conversion failed: {result.stderr}") 
            return None 
    except Exception as e: 
        print(f"PDF conversion error: {e}") 
        return None 
 
def process_job_application(job, profile, gemini_engine): 
    master_cv_path = "master_cv.docx" 
    if not os.path.exists(master_cv_path): 
        return None, None, "master_cv.docx not found in project folder" 
 
    cv_text, _ = load_master_cv(master_cv_path) 
    cv_summary = extract_cv_summary(cv_text) 
 
    tailored_summary = gemini_engine.tailor_cv(cv_summary, job['description'], profile) 
    cover_letter = gemini_engine.generate_cover_letter(job['company'], job['title'], job['description'], profile) 
 
    cv_path, cl_path, folder = create_tailored_cv( 
        job['id'], job['company'], job['title'], 
        tailored_summary, cover_letter, master_cv_path 
    ) 
 
    return cv_path, cl_path, folder 
 
def generate_application_package(job, profile, gemini_engine): 
    master_cv_path = "master_cv.docx" 
    if not os.path.exists(master_cv_path): 
        return None, None, None, "master_cv.docx not found" 
 
    print(f"\n🔄 Generating application package for {job['title']} at {job['company']}...") 
 
    # Load CV 
    cv_text, _ = load_master_cv(master_cv_path) 
    cv_summary = extract_cv_summary(cv_text) 
 
    # Generate tailored content 
    print("  → Tailoring CV summary with Gemini...") 
    tailored_summary = gemini_engine.tailor_cv(cv_summary, job['description'], profile) 
    if "error" in tailored_summary.lower() and "429" in tailored_summary: 
        return None, None, None, "AI quota exceeded — please try again in a few minutes" 
 
    print("  → Generating cover letter...") 
    cover_letter = gemini_engine.generate_cover_letter( 
        job['company'], job['title'], job['description'], profile 
    ) 
    if "error" in cover_letter.lower() and "429" in cover_letter: 
        return None, None, None, "AI quota exceeded — please try again in a few minutes" 
 
    # Create DOCX files 
    cv_path, cl_path, folder = create_tailored_cv( 
        job['id'], job['company'], job['title'], 
        tailored_summary, cover_letter, master_cv_path 
    ) 
 
    # Convert to PDF 
    print("  → Converting to PDF...") 
    cv_pdf = convert_docx_to_pdf(cv_path) 
    cl_pdf = convert_docx_to_pdf(cl_path) 
 
    # Save JD 
    jd_path = f"{folder}/JobDescription.txt" 
    with open(jd_path, 'w') as f: 
        f.write(f"Company: {job['company']}\n") 
        f.write(f"Role: {job['title']}\n") 
        f.write(f"Location: {job['location']}\n") 
        f.write(f"Source: {job['source']}\n") 
        f.write(f"URL: {job['url']}\n") 
        f.write(f"Date: {job['date_found']}\n") 
        f.write(f"Score: {job['score']}\n") 
        f.write(f"Score Reason: {job['score_reason']}\n\n") 
        f.write(f"FULL JOB DESCRIPTION:\n{job['description']}") 
 
    # Save tailored summary for reference 
    summary_path = f"{folder}/TailoredSummary.txt" 
    with open(summary_path, 'w') as f: 
        f.write(f"TAILORED CV SUMMARY:\n{tailored_summary}\n\n") 
        f.write(f"COVER LETTER:\n{cover_letter}") 
 
    print(f"✅ Application package ready: {folder}") 
    return cv_pdf or cv_path, cl_pdf or cl_path, folder, None 
 
if __name__ == "__main__": 
    print("CV Engine ready ✓") 
    cv_text, _ = load_master_cv("master_cv.docx") 
    summary = extract_cv_summary(cv_text) 
    print("Extracted summary preview:") 
    print(summary[:300])
