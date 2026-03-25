import io
import os
import json
import re
from docx import Document
from docx.shared import Pt
from datetime import date


def extract_cv_summary(cv_bytes):
    try:
        doc = Document(io.BytesIO(cv_bytes))
        full_text = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        summary_lines = []
        capture = False
        for line in full_text:
            if (
                "EXECUTIVE PROFILE" in line.upper()
                or "EXECUTIVE SUMMARY" in line.upper()
                or "PROFESSIONAL SUMMARY" in line.upper()
                or "CAREER SUMMARY" in line.upper()
                or "PROFILE" in line.upper()
                or "SUMMARY" in line.upper()
            ):
                capture = True
                continue
            if capture and line:
                if any(
                    s in line.upper()
                    for s in ["EXPERIENCE", "COMPETENCIES", "EDUCATION", "SKILLS", "CORE"]
                ):
                    break
                summary_lines.append(line)
            if len(summary_lines) >= 4:
                break
        return " ".join(summary_lines)
    except Exception:
        return ""


def tailor_cv_summary(cv_summary, job, profile, groq_key):
    if not groq_key or groq_key == "test_mode":
        years = profile.get("years_experience", 10)
        markets = profile.get("experience_markets", [])
        markets_text = ", ".join(markets) if markets else "Europe, APAC and global programmes"
        return "\n".join(
            [
                f"• {years}+ years leading recruitment/RPO delivery across {markets_text}.",
                f"• Strong track record hiring for senior roles similar to {job['title']} at {job['company']}.",
                "• Proven ability to manage stakeholders, drive process improvements and hit hiring SLAs.",
                "• Open to relocation with immediate or short‑notice availability.",
            ]
        )

    try:
        from groq import Groq

        client = Groq(api_key=groq_key)
        prompt = f"""Rewrite this CV summary into 3–5 sharp bullet points for a tailored CV. 

Return ONLY the bullet lines, each on its own line starting with "• ". No intro text, no heading.

 CANDIDATE: {profile.get('name')}, {profile.get('years_experience')} years experience 
 MARKETS: {profile.get('experience_markets',[])} 
 TARGET JOB: {job['title']} at {job['company']} in {job['location']} 
 JD: {job.get('description','')[:600]} 
 
 CURRENT SUMMARY: {cv_summary} 

 Rules: 
 - 3–5 concise bullet points only 
 - Start each line with "• " 
 - Mirror JD keywords naturally 
 - Include concrete metrics where possible 
 - End one bullet with relocation/availability statement 
 - Do not include company client names"""

        r = client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.4,
        )
        return r.choices[0].message.content.strip()
    except Exception:
        years = profile.get("years_experience", 10)
        markets = profile.get("experience_markets", [])
        markets_text = ", ".join(markets) if markets else "Europe, APAC and global programmes"
        return "\n".join(
            [
                f"• {years}+ years leading recruitment/RPO delivery across {markets_text}.",
                f"• Strong track record hiring for senior roles similar to {job['title']} at {job['company']}.",
                "• Proven ability to manage stakeholders, drive process improvements and hit hiring SLAs.",
                "• Open to relocation with immediate or short‑notice availability.",
            ]
        )


def generate_cover_letter(job, profile, groq_key):
    if not groq_key or groq_key == "test_mode":
        years = profile.get("years_experience", 5)
        markets = profile.get("experience_markets", [])
        markets_text = ", ".join(markets) if markets else "different markets and environments"
        skills = profile.get("skills") or []
        skills_text = ", ".join(skills[:6]) if skills else "relevant roles"
        return f"""Dear Hiring Team, 

 I am excited to apply for the {job['title']} position at {job['company']}. Based on the role description, I believe my background and experience are a strong match for your requirements. 
 
 With {years}+ years of experience across {markets_text}, I have developed deep expertise in {skills_text} and a track record of delivering measurable results. I am confident that this experience will allow me to contribute quickly and effectively in this role. 
 
 I am highly motivated by the opportunity to join {job['company']} and would welcome the chance to discuss my application further. 
 
 Best regards, 
 {profile.get('name','')} 
 {profile.get('email','')} 
 {profile.get('phone','')}"""

    try:
        from groq import Groq

        client = Groq(api_key=groq_key)
        prompt = f"""Write a 3-paragraph cover letter for {profile.get('name')} applying to {job['title']} at {job['company']}. 

 Profile: {profile.get('years_experience')} years, markets: {profile.get('experience_markets',[])}, relocate: {profile.get('relocate',True)} 
 JD: {job.get('description','')[:500]} 
 
 Rules: 
 - Para 1: Why this company and role 
 - Para 2: Why candidate fits with real metrics 
 - Para 3: Relocation commitment and close 
 - No client names 
 - Professional warm tone 
 - Return only letter body, no greeting/sign-off"""

        r = client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.4,
        )
        body = r.choices[0].message.content.strip()
        return f"""Dear Hiring Team, 

 {body} 
 
 Best regards, 
 {profile.get('name','')} 
 {profile.get('email','')} 
 {profile.get('phone','')}"""
    except Exception:
        years = profile.get("years_experience", 5)
        markets = profile.get("experience_markets", [])
        markets_text = ", ".join(markets) if markets else "different markets and environments"
        skills = profile.get("skills") or []
        skills_text = ", ".join(skills[:6]) if skills else "relevant roles"
        return f"""Dear Hiring Team, 

 I am excited to apply for the {job['title']} position at {job['company']}. Based on the role description, I believe my background and experience are a strong match for your requirements. 
 
 With {years}+ years of experience across {markets_text}, I have developed deep expertise in {skills_text} and a track record of delivering measurable results. I am confident that this experience will allow me to contribute quickly and effectively in this role. 
 
 I am highly motivated by the opportunity to join {job['company']} and would welcome the chance to discuss my application further. 
 
 Best regards, 
 {profile.get('name','')} 
 {profile.get('email','')} 
 {profile.get('phone','')}"""


def create_tailored_cv_bytes(cv_bytes, tailored_summary, job):
    try:
        doc = Document(io.BytesIO(cv_bytes))
        profile_found = False
        replaced = False
        for para in doc.paragraphs:
            if (
                "EXECUTIVE PROFILE" in para.text.upper()
                or "PROFILE" in para.text.upper()
                or "SUMMARY" in para.text.upper()
            ):
                profile_found = True
                continue
            if profile_found and not replaced and para.text.strip():
                if any(
                    s in para.text.upper()
                    for s in ["EXPERIENCE", "COMPETENCIES", "EDUCATION", "SKILLS", "CORE"]
                ):
                    break
                for run in para.runs:
                    run.text = ""
                if para.runs:
                    para.runs[0].text = tailored_summary
                else:
                    para.add_run(tailored_summary)
                replaced = True
        if not replaced:
            for para in doc.paragraphs:
                if para.text.strip():
                    for run in para.runs:
                        run.text = ""
                    if para.runs:
                        para.runs[0].text = tailored_summary
                    else:
                        para.add_run(tailored_summary)
                    break
        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        return output.getvalue()
    except Exception:
        return cv_bytes


def extract_basic_profile_from_cv(cv_bytes):
    profile = {
        "name": "",
        "email": "",
        "phone": "",
        "location": "",
        "years_experience": 0,
    }
    try:
        doc = Document(io.BytesIO(cv_bytes))
        lines = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        text = "\n".join(lines)
    except Exception:
        return profile
    m = re.search(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}", text)
    if m:
        profile["email"] = m.group(0)
    m = re.search(r"(\+?\d[\d\s\-\(\)]{7,})", text)
    if m:
        profile["phone"] = m.group(1).strip()
    for line in lines[:10]:
        up = line.upper()
        if "CURRICULUM" in up or "RESUME" in up:
            continue
        if "@" in line:
            continue
        if any(c.isalpha() for c in line):
            profile["name"] = line.strip()
            break
    for line in lines[:15]:
        if line == profile["name"]:
            continue
        if "@" in line:
            continue
        if sum(1 for c in line if c.isdigit()) > 4:
            continue
        if "(" in line or ")" in line:
            continue
        words = line.replace(",", " ").split()
        if len(words) == 0 or len(words) > 6:
            continue
        lower = line.lower()
        known_places = [
            "india",
            "netherlands",
            "germany",
            "france",
            "spain",
            "italy",
            "uk",
            "united kingdom",
            "usa",
            "united states",
            "canada",
            "dubai",
            "uae",
            "singapore",
            "australia",
        ]
        if not any(k in lower for k in known_places) and not re.search(
            r"[A-Za-z]+\s*,\s*[A-Za-z]+", line
        ):
            continue
        profile["location"] = line.strip()
        break
    m = re.search(r"(\d{1,2})\+?\s+years", text, flags=re.IGNORECASE)
    if m:
        try:
            years = int(m.group(1))
            if 0 < years <= 40:
                profile["years_experience"] = years
        except Exception:
            pass
    return profile


def infer_target_roles_from_cv(cv_bytes, groq_key=None):
    summary = extract_cv_summary(cv_bytes)
    text = summary
    if not text:
        try:
            doc = Document(io.BytesIO(cv_bytes))
            lines = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            text = " ".join(lines[:40])
        except Exception:
            text = ""
    if not text:
        return []
    roles = []
    if groq_key and groq_key != "test_mode":
        try:
            from groq import Groq

            client = Groq(api_key=groq_key)
            prompt = f"""From this CV text, list 3-5 realistic target job titles for this candidate.
Return ONLY a valid JSON list of strings, no extra text.

CV:
{text[:2000]}"""
            r = client.chat.completions.create(
                model="llama-3.3-70b-versatile",
                messages=[{"role": "user", "content": prompt}],
                temperature=0.2,
            )
            out = r.choices[0].message.content.strip()
            if "```" in out:
                out = out.split("```")[1]
                if out.startswith("json"):
                    out = out[4:]
            data = json.loads(out.strip())
            roles = [str(x).strip() for x in data if str(x).strip()]
        except Exception:
            roles = []
    if roles:
        return roles
    profile_hint = extract_basic_profile_from_cv(cv_bytes)
    name_hint = (profile_hint.get("name") or "").upper()
    core_role_words = [
        "recruitment",
        "talent acquisition",
        "talent",
        "people",
        "human resources",
        "hr",
        "operations",
        "project manager",
        "product manager",
        "program manager",
        "delivery",
        "staffing",
        "sourcing",
        "resourcing",
        "consultant",
        "partner",
        "specialist",
        "lead",
        "leader",
        "head",
        "director",
        "manager",
        "vp",
        "vice president",
        "chief",
    ]
    seniority_words = [
        "head",
        "director",
        "senior",
        "lead",
        "vp",
        "vice president",
        "chief",
        "manager",
    ]
    location_words = [
        "india",
        "europe",
        "spain",
        "netherlands",
        "germany",
        "france",
        "italy",
        "uk",
        "united kingdom",
        "usa",
        "united states",
        "canada",
        "dubai",
        "uae",
        "singapore",
        "australia",
    ]
    def extract_role_only(fragment):
        low = fragment.lower()
        positions = []
        for phrase in core_role_words + seniority_words:
            idx = low.find(phrase)
            if idx != -1:
                positions.append(idx)
        if not positions:
            return ""
        start = min(positions)
        role_part = fragment[start:].strip(" ,;-")
        if len(role_part) < 4 or len(role_part) > 80:
            return ""
        return role_part

    text_norm = text.replace("|", " ").replace(",", " ")
    parts = re.split(r"[•\n;\.]", text_norm)
    candidates = []
    for part in parts:
        s = part.strip()
        if not s:
            continue
        up = s.upper()
        if name_hint and up.startswith(name_hint):
            continue
        if len(s) < 4 or len(s) > 90:
            continue
        if "@" in s:
            continue
        low = s.lower()
        if not any(w in low for w in core_role_words):
            continue
        if any(l in low for l in location_words) and not any(
            w in low for w in seniority_words
        ):
            continue
        for cut in [" covering ", " across ", " spanning ", " within ", " across "]:
            idx = low.find(cut)
            if idx != -1:
                s = s[:idx].strip(",; -")
                low = s.lower()
        s_role = extract_role_only(s)
        if not s_role:
            continue
        if s_role and s_role not in candidates:
            candidates.append(s_role)
        if len(candidates) >= 5:
            break
    seen = set()
    roles_out = []
    for c in candidates:
        if c not in seen:
            seen.add(c)
            roles_out.append(c)
    return roles_out


def infer_markets_from_cv(cv_bytes):
    markets = []
    try:
        doc = Document(io.BytesIO(cv_bytes))
        lines = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        text = " ".join(lines[:80])
    except Exception:
        return markets
    lower = text.lower()
    mapping = {
        "India": ["india"],
        "Europe": ["europe", "european", "eu"],
        "UK": ["united kingdom", "uk", "london", "england", "scotland"],
        "North America": ["north america", "usa", "united states", "canada"],
        "Middle East": ["middle east", "gcc", "uae", "dubai", "saudi"],
        "APAC": ["apac", "asia pacific", "south east asia"],
        "Latin America": ["latin america", "latam", "brazil", "mexico"],
        "Global": ["global"],
        "Remote": ["remote"],
    }
    for label, keys in mapping.items():
        if any(k in lower for k in keys):
            markets.append(label)
    seen = set()
    unique = []
    for mkt in markets:
        if mkt not in seen:
            seen.add(mkt)
            unique.append(mkt)
    return unique


def infer_industries_from_cv(cv_bytes):
    industries = []
    try:
        doc = Document(io.BytesIO(cv_bytes))
        lines = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        text = " ".join(lines[:80]).lower()
    except Exception:
        return industries
    mapping = {
        "Life Sciences": ["pharma", "pharmaceutical", "life sciences", "biotech"],
        "Tech": ["saas", "software", "technology", "tech", "digital"],
        "Finance": ["bank", "financial", "finance", "fintech", "investment"],
        "Consulting": ["consulting", "consultant"],
        "Retail": ["retail", "ecommerce", "e-commerce"],
        "Manufacturing": ["manufacturing", "plant", "factory"],
        "Healthcare": ["hospital", "healthcare", "clinic"],
        "Public Sector": ["government", "public sector"],
    }
    for label, keys in mapping.items():
        if any(k in text for k in keys):
            industries.append(label)
    seen = set()
    unique = []
    for ind in industries:
        if ind not in seen:
            seen.add(ind)
            unique.append(ind)
    return unique


if __name__ == "__main__":
    print("CV public engine ready ✓")
